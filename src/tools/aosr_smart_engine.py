"""
Умный генератор АОСР (Smart Engine).
Генерирует документ программно без использования внешних шаблонов docxtpl.
Обеспечивает 100% точность позиционирования данных.
"""
import os
from datetime import datetime
from typing import Type
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from crewai.tools import BaseTool
from pydantic import BaseModel, Field

class SmartAOSRSchema(BaseModel):
    act_number: str = Field(..., description="Номер акта")
    object_name: str = Field(..., description="Объект и адрес")
    customer_fullname: str = Field(..., description="Заказчик (полное наименование, ИНН, ОГРН, адрес)")
    contract_info: str = Field(..., description="Реквизиты контракта")
    work_name: str = Field(..., description="Наименование выполненных работ")
    start_date: str = Field(..., description="Дата начала")
    end_date: str = Field(..., description="Дата окончания")
    materials_list: str = Field(..., description="Материалы и сертификаты")
    docs_list: str = Field(..., description="Исполнительные схемы, протоколы")
    next_work: str = Field(..., description="Последующие работы")

class SmartAOSRTool(BaseTool):
    name: str = "smart_aosr_generator"
    description: str = "Генерирует юридически верный АОСР (Минстрой №344/пр) полностью программным способом."
    args_schema: Type[BaseModel] = SmartAOSRSchema

    def _run(self, **kwargs) -> str:
        try:
            doc = Document()
            
            # Настройка шрифтов
            style = doc.styles['Normal']
            style.font.name = 'Times New Roman'
            style.font.size = Pt(11)

            # 1. Шапка (Приложение №3)
            header = doc.add_paragraph()
            header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = header.add_run("Приложение № 3\nк приказу Министерства строительства\nи жилищно-коммунального хозяйства\nРоссийской Федерации\nот 16 мая 2023 г. № 344/пр")
            run.font.size = Pt(9)
            run.italic = True

            # 2. Заголовок
            title = doc.add_paragraph()
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_t = title.add_run("\n\nАКТ\nосвидетельствования скрытых работ")
            run_t.bold = True
            run_t.font.size = Pt(12)
            
            meta = doc.add_paragraph()
            meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
            meta.add_run(f"№ {kwargs['act_number']} от {datetime.now().strftime('%d.%m.%Y')}")

            # 3. Основные блоки (Таблицы или параграфы)
            def add_field(label, value):
                p = doc.add_paragraph()
                p.add_run(f"{label}: ").bold = True
                p.add_run(str(value))

            add_field("Объект капитального строительства", kwargs['object_name'])
            add_field("Застройщик (Технический заказчик)", kwargs['customer_fullname'])
            
            # Жестко зашитый Подрядчик из базы (не меняется агентом)
            contractor_fixed = (
                "ООО «СТРОЙСТАНДАРТ», ИНН 6163227552, ОГРН 1226100028912. "
                "Адрес: 344068, Ростовская обл., г. Ростов-на-Дону, пр-кт Михаила Нагибина, зд. 40Б, пом. 4. "
                "Генеральный директор: ТАМРАЗОВ АРТЕМ СЕРОБОВИЧ"
            )
            add_field("Лицо, осуществляющее строительство", contractor_fixed)
            
            doc.add_paragraph("\nПроизвели осмотр работ и составили настоящий акт:")
            
            add_field("1. К освидетельствованию предъявлены работы", kwargs['work_name'])
            add_field("2. Работы выполнены согласно контракту", kwargs['contract_info'])
            add_field("3. Применены материалы", kwargs['materials_list'])
            add_field("4. Предъявлены документы", kwargs['docs_list'])
            add_field("5. Сроки", f"с {kwargs['start_date']} по {kwargs['end_date']}")
            
            # 4. Решение
            res_p = doc.add_paragraph("\nРЕШЕНИЕ:")
            res_p.add_run("\nРаботы выполнены в соответствии с требованиями. Разрешается производство последующих работ: ").italic = True
            res_p.add_run(kwargs['next_work']).bold = True

            # 5. Блок подписей (Таблица)
            doc.add_paragraph("\nПредставители:")
            table = doc.add_table(rows=3, cols=2)
            table.autofit = True
            
            # Подпись Директора (жесткая)
            table.cell(0, 0).text = "От Подрядчика:\nГенеральный директор"
            table.cell(0, 1).text = "________________ / Тамразов А.С."
            
            table.cell(1, 0).text = "От Заказчика:\nПредставитель"
            table.cell(1, 1).text = "________________ / ________________"
            
            table.cell(2, 0).text = "Производитель работ:"
            table.cell(2, 1).text = "________________ / (пусто)"

            # Сохранение
            output_dir = "output/aosr_smart"
            os.makedirs(output_dir, exist_ok=True)
            path = os.path.join(output_dir, f"SmartAOSR_{kwargs['act_number']}.docx")
            doc.save(path)
            
            return f"🚀 SMART-АКТ №{kwargs['act_number']} СОЗДАН. Путь: {os.path.abspath(path)}"

        except Exception as e:
            return f"❌ Ошибка SMART-генератора: {e}"
